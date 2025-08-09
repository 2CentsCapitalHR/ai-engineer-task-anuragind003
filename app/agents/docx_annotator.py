from __future__ import annotations

import os
from typing import List, Optional, Dict, Any
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT, CONTENT_TYPE as CT
from docx.opc.packuri import PackURI
import math
from datetime import datetime
from uuid import uuid4
from app.core.embeddings import embed_texts
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

from app.core.config import get_settings
from app.models.schemas import IssueItem
from app.core.llm import get_llm_client


def _find_paragraph_index(doc: Document, candidates: List[str]) -> Optional[int]:
    lowered_candidates = [c.lower() for c in candidates if c]
    if not lowered_candidates:
        return None
    for idx, p in enumerate(doc.paragraphs):
        txt = (p.text or "").lower()
        for c in lowered_candidates:
            if c and c[:30] in txt:  # fuzzy contains
                return idx
    return None


def _keyword_anchor(doc: Document, issue: IssueItem) -> Optional[int]:
    text = issue.issue.lower()
    keywords: List[str] = []
    if "jurisdiction" in text or "court" in text:
        keywords.extend(["jurisdiction", "court", "uae", "abu dhabi", "adgm"])
    if "address" in text:
        keywords.extend(["address", "registered office"]) 
    # Missing signature and missing ADGM presence are absences; keep unanchored
    if "signature" in text or ("adgm" in text and "not" in text):
        keywords = []

    if not keywords:
        return None

    best_idx: Optional[int] = None
    best_hits = 0
    for idx, p in enumerate(doc.paragraphs):
        pt = (p.text or "").lower()
        hits = sum(1 for k in keywords if k in pt)
        if hits > best_hits:
            best_hits = hits
            best_idx = idx
    return best_idx if best_hits > 0 else None


def _semantic_anchor(doc: Document, issue: IssueItem) -> Optional[int]:
    # Build a short query from issue text
    query = issue.issue
    if not query:
        return None
    paragraphs = [p.text or "" for p in doc.paragraphs]
    if not paragraphs:
        return None
    try:
        par_emb = embed_texts(paragraphs)
        q_emb = embed_texts([query])[0]
        # dot product works as cosine if normalized in embed_texts
        best_idx = None
        best_score = -1.0
        for i in range(len(par_emb)):
            score = float((par_emb[i] * q_emb).sum())
            if score > best_score:
                best_score = score
                best_idx = i
        # require a minimal similarity to avoid random matches
        if best_idx is not None and best_score >= 0.2:
            return int(best_idx)
    except Exception:
        return None
    return None


def _llm_anchor_map(doc: Document, issues: List[IssueItem]) -> Dict[int, int]:
    """Ask the LLM to pick paragraph indices for each issue. Returns mapping issue_idx->para_idx."""
    llm = get_llm_client()
    if llm is None or not issues:
        return {}
    # Build a compact list of paragraphs (index + first ~180 chars)
    paras = []
    for i, p in enumerate(doc.paragraphs):
        txt = (p.text or "").strip()
        if txt:
            paras.append({"idx": i, "text": txt[:180]})
        if len(paras) >= 60:
            break
    if not paras:
        return {}
    # Build prompt
    try:
        import json
        issues_payload = [
            {
                "idx": i,
                "issue": it.issue,
                "section": it.section,
            }
            for i, it in enumerate(issues)
        ]
        prompt = (
            "You are aligning review issues to document paragraphs. "
            "Given a list of paragraphs (idx,text) and a list of issues (idx,issue,section), "
            "return a JSON list of objects {issue_idx:int, paragraph_idx:int} selecting the most relevant paragraph for each issue. "
            "If no suitable paragraph exists, omit that issue. Do not add extra keys or text.\n\n"
            f"Paragraphs: {json.dumps(paras)}\n\nIssues: {json.dumps(issues_payload)}\n"
        )
        data = llm.generate_json_list(prompt)
        result: Dict[int, int] = {}
        for item in data:
            ii = item.get("issue_idx")
            pi = item.get("paragraph_idx")
            if isinstance(ii, int) and isinstance(pi, int):
                result[ii] = pi
        return result
    except Exception:
        return {}


def _get_comments_part(doc: Document):
    """Best-effort to get or create the comments part. Returns (comments_part, comments_element) or (None, None)."""
    try:
        # Newer python-docx may expose comments_part
        cp = getattr(doc.part, "comments_part", None)
        if cp is None:
            # Try private accessor
            cp = getattr(doc.part, "_comments_part", None)
        if cp is None:
            # Try to add a new comments part
            # Create /word/comments.xml and relate it to document
            partname = PackURI("/word/comments.xml")
            if not any(r.reltype == RT.COMMENTS for r in doc.part.rels.values()):
                comments_element = OxmlElement("w:comments")
                comments_element.set(qn("xmlns:w"), "http://schemas.openxmlformats.org/wordprocessingml/2006/main")
                # Attach part to package
                doc.part.package._parts[partname] = doc.part.package._package_reader._part_factory(partname, CT.WML_COMMENTS, comments_element)
                doc.part.relate_to(partname, RT.COMMENTS, is_external=False)
            # Try fetching again
            cp = getattr(doc.part, "comments_part", None) or getattr(doc.part, "_comments_part", None)
        comments_element = getattr(cp, "element", None) if cp is not None else None
        return cp, comments_element
    except Exception:
        return None, None


def _add_word_comment(paragraph, comment_text: str, author: str = "ADGM Agent", initials: str = "AA") -> bool:
    try:
        doc = paragraph.part.document
    except Exception:
        # Fallback in some versions
        try:
            doc = paragraph._parent
        except Exception:
            return False
    cp, comments = _get_comments_part(doc)
    if comments is None:
        return False
    # determine next comment id
    try:
        existing = comments.xpath(".//w:comment", namespaces=comments.nsmap)
        next_id = 0
        for c in existing:
            cid = c.get(qn("w:id"))
            if cid is not None:
                next_id = max(next_id, int(cid))
        next_id += 1
    except Exception:
        next_id = 0

    # create comment element
    c = OxmlElement("w:comment")
    c.set(qn("w:id"), str(next_id))
    c.set(qn("w:author"), author)
    c.set(qn("w:initials"), initials)
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = comment_text
    r.append(t)
    p.append(r)
    c.append(p)
    comments.append(c)

    # add range start/end + reference in the paragraph
    p_el = paragraph._p
    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), str(next_id))
    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), str(next_id))
    ref_r = OxmlElement("w:r")
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), str(next_id))
    ref_r.append(ref)

    p_el.insert(0, start)
    p_el.append(end)
    p_el.append(ref_r)
    return True
    


def _add_inline_comment(paragraph, note_text: str) -> None:
    # Highlight paragraph to draw attention
    for run in paragraph.runs:
        try:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        except Exception:
            pass
    # Append a colored inline note
    note = paragraph.add_run(f" [Comment: {note_text}] ")
    note.font.color.rgb = RGBColor(229, 57, 53)  # red
    note.italic = True


def _short_citation(ev) -> str:
    snippet = getattr(ev, "snippet", "") or ""
    ref_id = getattr(ev, "ref_id", "") or ""
    # Try to extract a compact law/article reference from the snippet
    import re
    m = re.search(r"(ADGM[^\n\r,]*?Regulations\s*\d{4}).{0,40}?(Article|Art\.)\s*(\d+[A-Za-z]?)", snippet, flags=re.IGNORECASE)
    if m:
        law = m.group(1).strip()
        art = m.group(3).strip()
        return f"{law}, Art. {art}"
    # Fallback to ref_id
    return ref_id


def annotate_docx(original_path: str, issues: List[IssueItem]) -> str:
    settings = get_settings()
    os.makedirs(settings.output_dir, exist_ok=True)

    doc = Document(original_path)

    # Group issues per paragraph to avoid multiple long inline notes
    paragraph_to_issues: Dict[int, List[IssueItem]] = {}
    unanchored: List[IssueItem] = []

    for issue in issues:
        candidates: List[str] = []
        if issue.evidence and issue.evidence[0].snippet:
            candidates.append(issue.evidence[0].snippet)
        if issue.issue:
            candidates.append(issue.issue)
        if issue.section:
            candidates.append(issue.section)
        p_idx = _find_paragraph_index(doc, candidates)
        if p_idx is None:
            # Try keyword-based anchor
            p_idx = _keyword_anchor(doc, issue)
        if p_idx is None:
            # Try semantic anchor (embeddings)
            p_idx = _semantic_anchor(doc, issue)
        if p_idx is not None:
            paragraph_to_issues.setdefault(p_idx, []).append(issue)
        else:
            unanchored.append(issue)

    # LLM anchoring to refine paragraph targets
    llm_map = _llm_anchor_map(doc, issues)

    # Insert one concise inline note per paragraph
    for p_idx, iss_list in sorted(paragraph_to_issues.items()):
        snippets: List[str] = []
        for i, it in enumerate(iss_list, 1):
            cite = _short_citation(it.evidence[0]) if it.evidence else ""
            sev = f"[{it.severity}] " if it.severity else ""
            sug = f" Suggestion: {it.suggestion}" if it.suggestion else ""
            # Keep inline note concise; full details remain in the Issues section
            part = f"{i}. {sev}{it.issue}" + (f" (Per {cite})" if cite else "") + sug
            snippets.append(part)
        note_text = "; ".join(snippets)
        # Prefer LLM-selected paragraph for the first issue if available
        target_idx = llm_map.get(iss_list[0] and issues.index(iss_list[0]), p_idx)
        # Add both a Word comment (Review pane) and a short inline highlight note
        _add_word_comment(doc.paragraphs[target_idx], note_text)
        _add_inline_comment(doc.paragraphs[target_idx], note_text)

    # For unanchored items, add visible notes at end
    if unanchored:
        doc.add_paragraph()
        doc.add_heading("Unanchored Comments", level=2)
        for it in unanchored:
            cite = _short_citation(it.evidence[0]) if it.evidence else ""
            sev = f"[{it.severity}] " if it.severity else ""
            sug = f" Suggestion: {it.suggestion}" if it.suggestion else ""
            line = f"- {sev}{it.issue}" + (f" (Per {cite})" if cite else "") + sug
            doc.add_paragraph(line)

    # Append consolidated Issues section
    doc.add_page_break()
    doc.add_heading("Issues Found (Automated Review)", level=1)
    if not issues:
        doc.add_paragraph("No issues were detected.")
    else:
        for idx, issue in enumerate(issues, start=1):
            p = doc.add_paragraph()
            run_head = p.add_run(f"{idx}. [{issue.severity}] ")
            run_head.bold = True
            p.add_run(f"{issue.issue} ")
            # concise citation inline
            if issue.evidence:
                cite = _short_citation(issue.evidence[0])
                if cite:
                    p.add_run(f"(Per {cite})")
            if issue.suggestion:
                doc.add_paragraph(f"Suggestion: {issue.suggestion}")
            if issue.evidence:
                for ev in issue.evidence:
                    src = f" (source: {ev.source_url})" if getattr(ev, "source_url", None) else ""
                    doc.add_paragraph(f"Citation {ev.ref_id}: {ev.snippet}{src}")

    base = os.path.basename(original_path)
    name, ext = os.path.splitext(base)
    out_path = os.path.join(settings.output_dir, f"{name}_reviewed{ext}")
    try:
        doc.save(out_path)
    except PermissionError:
        # Fallback to a unique filename to avoid Windows file-lock issues during overwrite
        unique_suffix = datetime.now().strftime("%Y%m%d-%H%M%S") + "-" + uuid4().hex[:6]
        out_path = os.path.join(settings.output_dir, f"{name}_reviewed_{unique_suffix}{ext}")
        doc.save(out_path)
    return out_path
