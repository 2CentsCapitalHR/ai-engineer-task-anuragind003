from __future__ import annotations

import os
from docx import Document


ROOT = os.path.dirname(os.path.dirname(__file__))
UPLOAD_DIR = os.path.join(ROOT, "data", "uploads")


def write_doc(path: str, title: str, paragraphs: list[str]) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(path)


def main() -> None:
    os.makedirs(UPLOAD_DIR, exist_ok=True)

    # 1) Articles of Association with non-ADGM jurisdiction and missing signature
    aoa_path = os.path.join(UPLOAD_DIR, "sample_articles_of_association.docx")
    write_doc(
        aoa_path,
        "Articles of Association of SampleCo LTD",
        [
            "1. Name and incorporation: SampleCo LTD is incorporated as a limited company.",
            "2. Jurisdiction: Any dispute shall be subject to the UAE Federal Courts.",
            "3. Governance: The board may adopt by-laws as necessary.",
            "4. Meetings: General meetings shall be called by the directors.",
            # Intentionally avoid the word 'signature' to trigger the signatory issue
            # Intentionally omit any explicit 'ADGM' reference to trigger ADGM presence suggestion
        ],
    )

    # 2) Memorandum of Association (contains ADGM mention)
    moa_path = os.path.join(UPLOAD_DIR, "sample_memorandum_of_association.docx")
    write_doc(
        moa_path,
        "Memorandum of Association of SampleCo LTD",
        [
            "The company is established under the Abu Dhabi Global Market (ADGM) Companies Regulations.",
            "Objects: To conduct lawful business activities permitted by ADGM.",
            "Subscribers: The undersigned agree to subscribe for shares.",
            "Signed by: Authorized Person",
        ],
    )

    # 3) Board Resolution (neutral)
    br_path = os.path.join(UPLOAD_DIR, "sample_board_resolution.docx")
    write_doc(
        br_path,
        "Board Resolution of SampleCo LTD",
        [
            "RESOLVED THAT the Company proceed with incorporation formalities and appoint authorized signatories.",
            "FURTHER RESOLVED THAT the directors are authorized to submit applications as required.",
            "Signed by: Director",
        ],
    )

    print("Generated sample files:")
    for p in [aoa_path, moa_path, br_path]:
        print(" -", p)


if __name__ == "__main__":
    main()


